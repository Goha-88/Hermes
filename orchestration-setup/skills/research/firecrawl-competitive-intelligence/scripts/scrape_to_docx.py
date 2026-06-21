#!/usr/bin/env python3
"""
Convert scraped website content to a structured DOCX report.

Usage:
  python3 scrape_to_docx.py --url https://target.kz --output report.docx

Or use from within Hermes by passing scraped markdown content as args.
The script accepts scraped pages as --page "Title::URL::content" arguments.
"""

import argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(pages: list[dict], output_path: str, site_map: list[str] = None):
    """
    pages: list of dicts with keys: title, url, content (markdown)
    site_map: optional list of "name: url" strings
    """
    doc = Document()

    title = doc.add_heading(f'Website Content Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if pages:
        main_url = pages[0].get('url', '')
        doc.add_paragraph(f'Source: {main_url}\nPages scraped: {len(pages)}')

    for i, page in enumerate(pages):
        doc.add_page_break()
        doc.add_heading(f'{i+1}. {page["title"]}', 1)
        doc.add_paragraph(f'URL: {page["url"]}')

        content = page.get('content', '')
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Convert markdown headings to docx headings
            if line.startswith('# '):
                doc.add_heading(line[2:], 2)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 3)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 4)
            elif line.startswith('http'):
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)

    # Sitemap section
    if site_map:
        doc.add_page_break()
        doc.add_heading('Site Map (All URLs)', 1)
        for entry in site_map:
            doc.add_paragraph(entry)

    doc.save(output_path)
    print(f"DONE: {output_path}")


def page_from_args(arg_str: str) -> dict:
    """Parse 'Title::URL::content' format"""
    parts = arg_str.split('::', 2)
    return {
        'title': parts[0] if len(parts) > 0 else 'Untitled',
        'url': parts[1] if len(parts) > 1 else '',
        'content': parts[2] if len(parts) > 2 else '',
    }


def main():
    parser = argparse.ArgumentParser(description='Build DOCX from scraped pages')
    parser.add_argument('--page', action='append', help='Page as Title::URL::content (repeatable)')
    parser.add_argument('--sitemap', action='append', help='Sitemap entry as name: url (repeatable)')
    parser.add_argument('--output', '-o', default='report.docx', help='Output path')
    args = parser.parse_args()

    pages = [page_from_args(p) for p in (args.page or [])]
    site_map = args.sitemap or []

    if not pages:
        print("No pages provided. Use --page 'Title::URL::content'")
        return

    build_docx(pages, args.output, site_map)


if __name__ == '__main__':
    main()
